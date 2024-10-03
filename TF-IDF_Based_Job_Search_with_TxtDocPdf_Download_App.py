# Import necessary libraries
import pandas as pd
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import io
from docx import Document  # For generating Word document

# Load the job content data from your GitHub repository
url_csv = "https://raw.githubusercontent.com/tayler-erbe/Job_Posting_Generator_App/main/job_content.csv"
job_content_df = pd.read_csv(url_csv)

# Preprocess the 'Combine_String' column, filling NaN values with empty strings
combine_strings = job_content_df['Combine_String'].fillna('')

# Initialize the TF-IDF vectorizer
tfidf_vectorizer = TfidfVectorizer(stop_words='english')

# Fit and transform the job content using the TF-IDF vectorizer
tfidf_matrix = tfidf_vectorizer.fit_transform(combine_strings)

# Function to retrieve similar jobs using TF-IDF and cosine similarity
def retrieve_similar_jobs_tfidf(new_job_summary, tfidf_matrix, combine_strings, top_k=5):
    # Transform the new job query using the same TF-IDF vectorizer
    query_tfidf = tfidf_vectorizer.transform([new_job_summary])
    
    # Calculate cosine similarity between the query and job content
    cosine_similarities = cosine_similarity(query_tfidf, tfidf_matrix).flatten()
    
    # Get the top k job indices based on similarity scores
    top_k_indices = cosine_similarities.argsort()[-top_k:][::-1]
    
    # Retrieve the top k most similar job descriptions
    similar_jobs = job_content_df.iloc[top_k_indices]
    return similar_jobs

# Function to create a text file for download
def create_text_file(selected_job):
    output_text = f"""Job Title: {selected_job['Class Title']}\n
Job Duties: {selected_job['Job Duties']}\n
Minimum Qualifications: {selected_job['MAQ']}\n
Levels of Work: {selected_job['Levels Of Work']}\n
Knowledge, Skills, Abilities (KSA): {selected_job['KSA']}\n"""
    return output_text

# Function to create a Word document for download
def create_word_doc(selected_job):
    doc = Document()
    doc.add_heading(selected_job['Class Title'], 0)
    doc.add_paragraph(f"Job Duties: {selected_job['Job Duties']}")
    doc.add_paragraph(f"Minimum Qualifications: {selected_job['MAQ']}")
    doc.add_paragraph(f"Levels of Work: {selected_job['Levels Of Work']}")
    doc.add_paragraph(f"Knowledge, Skills, Abilities (KSA): {selected_job['KSA']}")
    return doc

# Streamlit App
st.title("TF-IDF Based Job Search App")

# User Input for Job Query
user_query = st.text_input("Enter a job query (e.g., 'I need an accountant with 4 years of experience'): ")

# Allow user to select the number of returned job results (k)
top_k = st.number_input("Enter number of job results to return", min_value=1, max_value=20, value=5)

# Retrieve jobs based on the user query
if user_query:
    # Retrieve jobs using the TF-IDF model
    similar_jobs_tfidf = retrieve_similar_jobs_tfidf(user_query, tfidf_matrix, combine_strings, top_k=top_k)
    
    # Display a selection box for the user to pick a job title
    selected_job_title = st.selectbox("Select a job title", similar_jobs_tfidf['Class Title'].tolist())
    
    # Display full job posting when a job title is selected
    if selected_job_title:
        selected_job = similar_jobs_tfidf[similar_jobs_tfidf['Class Title'] == selected_job_title].iloc[0]
        
        # Display job posting details
        st.write(f"### {selected_job['Class Title']}")
        st.write(f"**Job Duties:** {selected_job['Job Duties']}")
        st.write(f"**Minimum Qualifications:** {selected_job['MAQ']}")
        st.write(f"**Levels of Work:** {selected_job['Levels Of Work']}")
        st.write(f"**Knowledge, Skills, Abilities (KSA):** {selected_job['KSA']}")
        
        # Create a downloadable text file
        text_output = create_text_file(selected_job)
        st.download_button(
            label="Download as Text File",
            data=text_output,
            file_name=f"{selected_job['Class Title']}.txt",
            mime="text/plain"
        )
        
        # Create a downloadable Word document
        word_doc = create_word_doc(selected_job)
        word_buffer = io.BytesIO()
        word_doc.save(word_buffer)
        word_buffer.seek(0)
        st.download_button(
            label="Download as Word Document",
            data=word_buffer,
            file_name=f"{selected_job['Class Title']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
